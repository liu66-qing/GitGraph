from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "codegraph-demo-assets"
OUT = DOCS / "CodeGraph演示说明文档.docx"

ACCENT = "2563EB"
INK = "0F172A"
MUTED = "475569"
LIGHT = "EFF6FF"
SOFT = "F8FAFC"


MODULES = [
    {
        "title": "Graph Explorer（知识图谱浏览器）",
        "entry": "https://evo-graph.vercel.app/",
        "image": "01-graph-explorer.png",
        "caption": "图 1  Graph Explorer 知识图谱浏览器",
        "positioning": "用于展示 CodeGraph 已构建的实体关系网络，是系统知识图谱能力的主入口。",
        "features": [
            "左侧提供 5 个核心功能模块导航，便于在图谱、问答、文档、冲突和时序视图之间切换。",
            "中央区域基于 D3.js 展示力导向知识图谱，节点表示实体，连线表示实体之间的语义关系。",
            "不同颜色节点区分 Person、Organization、Product、Event、Location、Technology、Concept 等实体类型。",
            "顶部搜索框支持实体检索，右侧操作按钮支持缩放、重置等图谱浏览动作。",
            "图例面板对实体类型进行说明，便于评审快速理解图谱结构。",
        ],
        "acceptance": [
            "页面能够正常加载图谱可视化区域。",
            "实体节点、关系边和实体类型图例展示完整。",
            "搜索与视图控制入口清晰可见。",
        ],
    },
    {
        "title": "Query Console（智能问答控制台）",
        "entry": "https://evo-graph.vercel.app/query",
        "image": "02-query-console.png",
        "caption": "图 2  Query Console 智能问答控制台",
        "positioning": "用于验证 Agentic RAG 问答能力，体现系统在知识图谱和检索结果上的多跳推理与溯源回答。",
        "features": [
            "采用聊天式交互界面，用户可在底部输入自然语言问题。",
            "页面提供预设问题示例，覆盖事实查询、事件原因分析和对比分析等常见任务。",
            "回答区域用于展示最终答案、Reasoning Trace 推理轨迹和来源引用。",
            "后端 Agent 可按问题类型动态调用图查询、向量检索、时序查询和因果推理工具。",
            "问答流程以置信度为收敛目标，最多进行 5 轮迭代以提升回答可靠性。",
        ],
        "acceptance": [
            "输入框与示例问题入口可见。",
            "问答交互区域布局完整。",
            "文档说明中明确包含推理轨迹与溯源引用能力。",
        ],
    },
    {
        "title": "Document Ingestion（文档上传与知识演化）",
        "entry": "https://evo-graph.vercel.app/documents",
        "image": "03-document-ingestion.png",
        "caption": "图 3  Document Ingestion 文档上传与知识演化",
        "positioning": "用于展示从外部文档到知识图谱增量演化的处理链路，是系统知识更新能力的入口。",
        "features": [
            "支持拖拽上传和文件选择两种入口，面向 PDF、TXT、Markdown、HTML 等文档格式。",
            "上传后触发异步知识演化流水线，流程包括 Ingest、Extract、Resolve、Conflict Check、Merge。",
            "Extract 阶段通过 LLM 抽取实体和关系，Resolve 阶段进行实体消歧与去重。",
            "Conflict Check 阶段检测新增知识与历史事实之间的潜在冲突。",
            "Merge 阶段将确认后的新知识合并到图谱，并通过前端状态反馈展示处理进度。",
        ],
        "acceptance": [
            "上传区域、文件选择按钮和支持格式说明展示完整。",
            "五阶段 Evolution Pipeline 清晰可见。",
            "文档到图谱演化的核心流程描述准确。",
        ],
    },
    {
        "title": "Conflict Dashboard（知识冲突管理）",
        "entry": "https://evo-graph.vercel.app/conflicts",
        "image": "04-conflict-dashboard.png",
        "caption": "图 4  Conflict Dashboard 知识冲突管理",
        "positioning": "用于集中展示知识冲突检测结果，辅助用户对矛盾事实进行核查和裁决。",
        "features": [
            "页面左侧展示开放冲突列表，包含冲突类型、简要描述、检测时间和状态。",
            "系统覆盖 Temporal Overlap、Logical Contradiction、Source Disagreement 三类冲突。",
            "Temporal Overlap 用于识别同一角色在重叠时间窗口内被多个实体占据的情况。",
            "Logical Contradiction 用于识别互斥关系同时成立的情况。",
            "Source Disagreement 用于识别多个来源对同一事实给出不同描述的情况。",
        ],
        "acceptance": [
            "冲突列表能够展示冲突类型和状态。",
            "页面预留冲突详情查看区域。",
            "文档说明中明确三类冲突的含义和用途。",
        ],
    },
    {
        "title": "Timeline（时序演化回放）",
        "entry": "https://evo-graph.vercel.app/timeline",
        "image": "05-timeline.png",
        "caption": "图 5  Timeline 时序演化回放",
        "positioning": "用于展示知识图谱随时间演化的过程，支持回看历史状态和对比知识变化。",
        "features": [
            "顶部时间轴支持按时间点查看图谱状态，适合审查实体关系随时间变化的过程。",
            "事件列表记录实体创建、关系新增、关系过期、冲突检测等演化事件。",
            "图谱快照对比区域展示指定时间范围前后的实体和关系数量变化。",
            "每条关系可携带 valid_from 和 valid_to 时间窗口，从而支持历史时刻查询。",
            "适用于公司高管变动、产品迭代、政策法规更新等具有明显时间维度的知识场景。",
        ],
        "acceptance": [
            "时间轴、事件列表和快照对比区域展示完整。",
            "演化事件有清晰的类型标识和日期。",
            "文档说明中体现时序版本化能力。",
        ],
    },
]


def set_run_font(run, size=10.5, bold=False, color=INK):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, color=INK, size=9.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text="", size=10.5, bold=False, color=INK, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 12, bold=True, color=ACCENT if level == 1 else INK)
    return p


def add_callout(doc, title, lines, fill=LIGHT, trailing_space=True):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=9.8, bold=True, color=ACCENT)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(line)
        set_run_font(r, size=8.8, color=INK)
    if trailing_space:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc, items, size=9.2):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(item)
        set_run_font(r, size=size)


def image_for_doc(image_path):
    if image_path.suffix.lower() != ".png":
        return image_path
    jpg = image_path.with_suffix(".jpg")
    if not jpg.exists() or jpg.stat().st_mtime < image_path.stat().st_mtime:
        with Image.open(image_path) as img:
            img.convert("RGB").save(jpg, quality=92, optimize=True)
    return jpg


def add_image(doc, image_path, caption):
    if not image_path.exists():
        add_para(doc, f"[截图缺失：{image_path.name}]", size=9, color="B91C1C")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_for_doc(image_path)), width=Cm(14.8))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(caption)
    set_run_font(r, size=8.4, color=MUTED)


def add_info_table(doc):
    rows = [
        ("项目名称", "CodeGraph 实时知识图谱演化智能体"),
        ("在线 Demo", "https://evo-graph.vercel.app/"),
        ("GitHub", "https://github.com/liu66-qing/KG-RAG-Agent"),
        ("文档用途", "Demo 提交说明文档，用于展示系统页面、核心能力和验收关注点"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        cell_text(table.cell(i, 0), key, bold=True, color=ACCENT, size=9.5)
        cell_text(table.cell(i, 1), val, size=9.5)
        shade(table.cell(i, 0), "DBEAFE")


def add_capability_matrix(doc):
    rows = [
        ("知识构建", "文档上传后抽取实体与关系，并进行实体消歧、冲突检测和图谱合并。"),
        ("知识查询", "结合知识图谱、向量检索和时序查询，为自然语言问题生成可溯源答案。"),
        ("时序版本", "通过 valid_from / valid_to 描述关系有效期，支持历史状态回看和变化对比。"),
        ("冲突管理", "检测时序重叠、逻辑矛盾和来源分歧，并保留证据链供人工裁决。"),
        ("可视分析", "提供图谱浏览、问答控制台、文档流水线、冲突列表和演化时间轴五类页面。"),
    ]
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    cell_text(table.cell(0, 0), "能力项", bold=True, color=ACCENT)
    cell_text(table.cell(0, 1), "提交说明", bold=True, color=ACCENT)
    shade(table.cell(0, 0), "DBEAFE")
    shade(table.cell(0, 1), "DBEAFE")
    for idx, (key, val) in enumerate(rows, start=1):
        cell_text(table.cell(idx, 0), key, bold=True, color=MUTED)
        cell_text(table.cell(idx, 1), val)


def add_module(doc, module, index):
    add_heading(doc, f"{index}. {module['title']}", 1)
    meta = doc.add_table(rows=2, cols=2)
    meta.style = "Table Grid"
    cell_text(meta.cell(0, 0), "访问入口", bold=True, color=ACCENT)
    cell_text(meta.cell(0, 1), module["entry"])
    cell_text(meta.cell(1, 0), "功能定位", bold=True, color=ACCENT)
    cell_text(meta.cell(1, 1), module["positioning"])
    shade(meta.cell(0, 0), "DBEAFE")
    shade(meta.cell(1, 0), "DBEAFE")
    add_para(doc, after=2)
    add_image(doc, ASSETS / module["image"], module["caption"])
    add_heading(doc, "界面与功能说明", 2)
    add_bullets(doc, module["features"])
    add_callout(doc, "提交验收关注点", module["acceptance"], fill=SOFT)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("CodeGraph Demo 提交说明文档")
    set_run_font(r, size=23, bold=True, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("实时知识图谱演化智能体 | 页面截图、功能说明与验收要点")
    set_run_font(r, size=11, color=MUTED)

    add_info_table(doc)
    add_para(doc)
    add_heading(doc, "一、Demo 概述", 1)
    add_para(
        doc,
        "CodeGraph 是一个面向 Agentic RAG 场景的知识图谱演化系统。系统将文档解析、实体关系抽取、知识图谱构建、时序版本管理、冲突检测和可溯源问答整合到一个可演示的 Web 应用中。",
    )
    add_para(
        doc,
        "本提交文档基于在线 Demo 页面截图整理，重点说明各页面的功能定位、核心展示内容和评审验收时应关注的能力点。",
    )
    add_heading(doc, "二、核心能力矩阵", 1)
    add_capability_matrix(doc)
    add_heading(doc, "三、技术栈", 1)
    add_callout(
        doc,
        "系统实现",
        [
            "后端：Python + FastAPI，负责文档处理、Agent 推理编排和业务 API。",
            "知识层：Neo4j 知识图谱 + Qdrant 向量检索，支持图查询、语义检索和混合召回。",
            "模型层：DeepSeek LLM，用于实体关系抽取、推理和回答生成。",
            "前端：React + TypeScript + D3.js + TailwindCSS，提供图谱可视化和多页面交互。",
            "部署：Vercel 前端部署 + Render 后端部署，形成云端可访问 Demo。",
        ],
    )

    for index, module in enumerate(MODULES, start=1):
        if index > 1:
            doc.add_page_break()
        add_module(doc, module, index)

    add_heading(doc, "四、提交结论", 1)
    add_para(
        doc,
        "本 Demo 已覆盖知识图谱浏览、智能问答、文档上传、冲突管理和时序演化五个核心页面，可用于展示 CodeGraph 从文档到知识、从知识到推理、从推理到溯源的完整链路。核心创新点包括自适应推理循环、时序知识版本化、三类冲突检测和全链路溯源。",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
